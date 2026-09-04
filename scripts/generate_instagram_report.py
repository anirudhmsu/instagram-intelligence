import json
import statistics
import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def add_hyperlink(paragraph, label: str, url: str) -> None:
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_cell_text(cell, value: str, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(value)
    run.bold = bold
    run.font.size = Pt(9)


def build_report(source: Path, destination: Path) -> None:
    payload = json.loads(source.read_text(encoding="utf-8"))
    analysis = payload["analysis"]
    posts = analysis["posts"]
    trends = analysis["top_topics"]
    engagements = [post["engagement"] for post in posts]

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.color.rgb = RGBColor(31, 78, 121)

    title = document.add_heading("Instagram Engagement Data Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    selection_label = analysis.get("selection_method", "recent").replace("_", " ").title()
    subtitle = document.add_paragraph(f"@{analysis['username']} | {selection_label} post analysis")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True

    document.add_heading("1. Purpose and dataset scope", level=1)
    document.add_paragraph(
        "This report describes the structured JSON dataset collected for the public Instagram "
        f"profile @{analysis['username']}. The service inspected {analysis['scanned_post_count']} "
        f"posts and retained the {analysis['post_count']} posts with the highest available engagement. "
        "Engagement is calculated as likes plus comments supplied by Instagram."
    )
    document.add_paragraph(f"Research question: {payload.get('research_query') or 'Not specified'}")
    document.add_paragraph(f"Collection timestamp: {analysis['fetched_at']}")
    document.add_paragraph(f"Source JSON: {source.name}")

    document.add_heading("2. Executive summary", level=1)
    document.add_paragraph(payload["summary"])
    summary = document.add_table(rows=2, cols=4)
    summary.style = "Light Shading Accent 1"
    headers = ["Posts scanned", "Posts retained", "Top engagement", "Median engagement"]
    values = [
        str(analysis["scanned_post_count"]),
        str(len(posts)),
        f"{max(engagements):,}",
        f"{statistics.median(engagements):,.0f}",
    ]
    for index, header in enumerate(headers):
        set_cell_text(summary.rows[0].cells[index], header, True)
        set_cell_text(summary.rows[1].cells[index], values[index])

    document.add_heading("3. Selected posts", level=1)
    document.add_paragraph(
        f"The table uses the {selection_label.lower()} selection order. Post names are derived from the "
        "opening caption text because Instagram does not provide a separate post-title field."
    )
    table = document.add_table(rows=1, cols=6)
    table.style = "Light Shading Accent 1"
    for cell, header in zip(
        table.rows[0].cells,
        ["Rank", "Post name", "Date", "Topics", "Engagement", "Link"],
    ):
        set_cell_text(cell, header, True)
    for rank, post in enumerate(posts, 1):
        cells = table.add_row().cells
        date = post["published_at"].split("T", 1)[0]
        set_cell_text(cells[0], str(rank))
        set_cell_text(cells[1], post["post_name"])
        set_cell_text(cells[2], date)
        set_cell_text(cells[3], ", ".join(post["topics"]))
        set_cell_text(cells[4], f"{post['engagement']:,}")
        cells[5].text = ""
        add_hyperlink(cells[5].paragraphs[0], "Open post", post["url"])

    document.add_heading("4. Topic-level trends", level=1)
    document.add_paragraph(
        "A post may have more than one topic, so topic engagement totals can overlap and should not "
        "be added together. Momentum compares the newer and older halves of posts carrying that topic; "
        "a value above 1.0 indicates stronger engagement in the newer half."
    )
    topic_table = document.add_table(rows=1, cols=5)
    topic_table.style = "Light Shading Accent 1"
    for cell, header in zip(
        topic_table.rows[0].cells,
        ["Topic", "Posts", "Total engagement", "Average", "Momentum"],
    ):
        set_cell_text(cell, header, True)
    for trend in trends:
        cells = topic_table.add_row().cells
        values = [
            trend["topic"], str(trend["post_count"]), f"{trend['total_engagement']:,}",
            f"{trend['average_engagement']:,.2f}", f"{trend['momentum']:.2f}x",
        ]
        for cell, value in zip(cells, values):
            set_cell_text(cell, value)

    document.add_heading("5. Interpretation", level=1)
    top = posts[0]
    document.add_paragraph(
        f"The strongest individual post was “{top['post_name']}” with {top['engagement']:,} available "
        "interactions. Its engagement is substantially higher than the remainder of the selected set, "
        "so averages should be interpreted alongside the median."
    )
    document.add_paragraph(
        "The selected posts indicate that symptom-led hooks, practical health explanations, and concise "
        "question-based captions attract attention. These are descriptive observations from this profile, "
        "not causal evidence or medical recommendations."
    )

    document.add_heading("6. JSON structure and reuse", level=1)
    fields = [
        ("research_query", "The question attached to the collection request."),
        ("summary", "Automatically generated narrative overview."),
        ("analysis", "Profile-level metadata, topic trends, and detailed selected posts."),
        ("records", "Flat, dataframe-ready post records containing IDs, captions, dates, metrics and permalinks."),
        ("export_id / export_url", "Identifiers used by the backend to retrieve the saved export."),
    ]
    for name, meaning in fields:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(f"{name}: ").bold = True
        paragraph.add_run(meaning)

    document.add_heading("7. Data limitations", level=1)
    limitations = [
        "The ranking covers the scanned 100-post window, not necessarily the account’s complete history.",
        "Instagram may omit likes or comments in some payloads; unavailable values are stored as zero.",
        "Engagement is an absolute interaction count and is not normalized by follower count or reach.",
        "Topic labels are keyword-derived and should be reviewed before formal research use.",
        "Captions and metrics represent a point-in-time collection and may later change.",
    ]
    for item in limitations:
        document.add_paragraph(item, style="List Bullet")

    footer = section.footer.paragraphs[0]
    footer.text = "Generated from the Instagram Health Market Analyzer JSON export"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate a Word report from an ingestion JSON")
    parser.add_argument("source", type=Path)
    parser.add_argument("destination", type=Path)
    args = parser.parse_args()
    build_report(args.source, args.destination)
